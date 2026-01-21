"""
DOCX text extraction with structure preservation.
"""
from dataclasses import dataclass
from typing import List
from pathlib import Path
import re


@dataclass
class ExtractedSection:
    """Represents an extracted section from DOCX."""

    title: str
    content: str
    level: int  # Heading level (1, 2, 3, etc.)
    has_formula: bool
    has_table: bool


@dataclass
class ExtractedDocument:
    """Represents a fully extracted DOCX document."""

    filename: str
    chapter_number: int
    sections: List[ExtractedSection]
    full_text: str
    metadata: dict


class DocxExtractor:
    """
    Extract structured text from DOCX files.

    Uses python-docx to preserve document structure including:
    - Headings hierarchy
    - Paragraphs
    - Tables (converted to text)
    - Formula indicators (equations in text)
    """

    def extract(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text and structure from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractedDocument with sections and metadata
        """
        from docx import Document

        doc = Document(file_path)

        # Extract chapter number from filename
        chapter_num = self._extract_chapter_number(file_path.name)

        sections = []
        current_section_title = "Introduction"
        current_section_content = []
        current_level = 1

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            # Check if this is a heading
            if "Heading" in style_name or self._is_heading(text):
                # Save previous section
                if current_section_content:
                    content = "\n".join(current_section_content)
                    sections.append(
                        ExtractedSection(
                            title=current_section_title,
                            content=content,
                            level=current_level,
                            has_formula=self._has_formula(content),
                            has_table=False,
                        )
                    )

                # Start new section
                current_section_title = text
                current_section_content = []
                current_level = self._get_heading_level(style_name)
            else:
                current_section_content.append(text)

        # Don't forget the last section
        if current_section_content:
            content = "\n".join(current_section_content)
            sections.append(
                ExtractedSection(
                    title=current_section_title,
                    content=content,
                    level=current_level,
                    has_formula=self._has_formula(content),
                    has_table=False,
                )
            )

        # Extract tables
        for table in doc.tables:
            table_text = self._table_to_text(table)
            if table_text:
                sections.append(
                    ExtractedSection(
                        title="Table",
                        content=table_text,
                        level=3,
                        has_formula=False,
                        has_table=True,
                    )
                )

        full_text = "\n\n".join([s.content for s in sections])

        return ExtractedDocument(
            filename=file_path.name,
            chapter_number=chapter_num,
            sections=sections,
            full_text=full_text,
            metadata={
                "total_sections": len(sections),
                "has_tables": any(s.has_table for s in sections),
                "has_formulas": any(s.has_formula for s in sections),
                "word_count": len(full_text.split()),
            },
        )

    def _extract_chapter_number(self, filename: str) -> int:
        """Extract chapter number from filename like 'Chapter 1 - Notes (Final 1).docx'"""
        match = re.search(r"Chapter\s+(\d+)", filename, re.IGNORECASE)
        return int(match.group(1)) if match else 0

    def _is_heading(self, text: str) -> bool:
        """Heuristic to detect headings in text."""
        # Short text, all caps, or numbered section
        if len(text) < 100 and (
            text.isupper() or re.match(r"^\d+\.?\d*\s+", text)
        ):
            return True
        return False

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r"Heading\s*(\d+)", style_name)
        return int(match.group(1)) if match else 1

    def _has_formula(self, text: str) -> bool:
        """Check if text contains formulas."""
        formula_patterns = [
            r"[A-Za-z]\s*=\s*[A-Za-z0-9\*/\+\-\^]+",  # F = ma
            r"\d+\s*[x]\s*\d+",  # multiplication
            r"[^a-zA-Z0-9]",  # superscripts (Unicode)
        ]
        return any(re.search(p, text) for p in formula_patterns)

    def _table_to_text(self, table) -> str:
        """Convert a table to readable text format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows)
